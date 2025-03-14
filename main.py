from ingestion import read_xml, read_pdf, read_epub
from text_processing import clean_text
from summarization import summarize_text, analyze_comparative
from report_generation import generate_thesis, generate_report
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def process_book(file_path, file_type):
    """
    Process a book by:
      1. Reading the file content based on its file type (XML, PDF, or EPUB).
      2. Cleaning the extracted text to remove extra whitespace.
      3. Summarizing the cleaned text using the OpenAI API (with citations).
    
    Parameters:
      - file_path (str): Path to the book file.
      - file_type (str): The type of the file ('xml', 'pdf', or 'epub').

    Returns:
      - str: A summary of the book focused on social isolation.
    """
    # Read the book file according to its type.
    if file_type == "xml":
        raw_text = read_xml(file_path)
    elif file_type == "pdf":
        raw_text = read_pdf(file_path)
    elif file_type == "epub":
        raw_text = read_epub(file_path)
    else:
        raise ValueError("Unsupported file type")
    
    print("Finished reading file. Cleaning text...")
    # Clean the extracted text.
    cleaned = clean_text(raw_text)
    
    print("Summarizing entire text...")
    # Generate a summary of the cleaned text using the OpenAI API.
    summary = summarize_text(cleaned)
    return summary

def main():
    """
    Main function to execute the entire workflow:
      1. Define file paths for the three novels.
      2. Process each book to produce summaries.
      3. Analyze the summaries comparatively.
      4. Generate a thesis statement based on the summaries and comparative analysis.
      5. Generate a complete five-paragraph book report (with citations) using the thesis, summaries, and analysis.
      6. Export the final report as a DOCX file.
    """
    # Define file paths (assumes a "Readings" folder in the project root).
    xml_path = "Readings/The-Bell-Jar-1645639705._vanilla.xml"
    pdf_path = "Readings/the_stranger.pdf"
    epub_path = "Readings/franz-kafka_metamorphosis.epub"

    # Process each book to obtain its summary.
    print("Processing Metamorphosis (EPUB)...")
    metamorphosis_summary = process_book(epub_path, "epub")
    
    print("Processing The Bell Jar (XML)...")
    belljar_summary = process_book(xml_path, "xml")
    
    print("Processing The Stranger (PDF)...")
    stranger_summary = process_book(pdf_path, "pdf")
    
    # Combine the individual book summaries into a list.
    summaries = [belljar_summary, stranger_summary, metamorphosis_summary]
    print("Analyzing comparative themes...")
    # Analyze the summaries to extract thematic insights.
    comparative_analysis = analyze_comparative(summaries)
    
    print("Generating thesis statement...")
    # Generate a thesis statement based on the book summaries and their comparative analysis.
    thesis_statement = generate_thesis(summaries, comparative_analysis)
    print("Thesis generated:", thesis_statement)
    
    print("Generating final book report...")
    # Generate the final five-paragraph book report.
    book_summaries = [belljar_summary, stranger_summary, metamorphosis_summary]
    final_report = generate_report(thesis_statement, book_summaries, comparative_analysis)
    
    # Export the final report as a DOCX file.
    print("Exporting final report to DOCX...")
    document = Document()
    # Add a title to the document using a heading style.
    document.add_heading(final_report.split('\n\n')[0], 0)  # Assume title is the first line.
    
    # Split the report into paragraphs (separated by two newlines).
    paragraphs = final_report.split("\n\n")
    # Skip the title if already used.
    for para in paragraphs[1:]:
        p = document.add_paragraph(para)
        # Set first-line indent (e.g., 36 points ~ half inch).
        p.paragraph_format.first_line_indent = Pt(36)
        # Optionally, align text to justify.
        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    
    # Save the document.
    document.save("Final_Book_Report.docx")
    print("Book report generated and saved as Final_Book_Report.docx")

if __name__ == "__main__":
    main()
